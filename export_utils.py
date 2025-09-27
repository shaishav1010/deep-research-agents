"""Export utilities for generating PDF, Word, and XML documents"""
import io
from typing import Dict, Any
from datetime import datetime
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import dicttoxml
from models import ResearchReport, ReportSection


def generate_pdf(report: ResearchReport) -> bytes:
    """Generate PDF from research report"""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter,
                            rightMargin=72, leftMargin=72,
                            topMargin=72, bottomMargin=18)

    # Container for the 'Flowable' objects
    elements = []

    # Define styles
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor='#1E3A8A',
        spaceAfter=30,
        alignment=TA_CENTER
    )
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=16,
        textColor='#2563EB',
        spaceAfter=12,
        spaceBefore=12
    )
    body_style = ParagraphStyle(
        'CustomBody',
        parent=styles['BodyText'],
        fontSize=11,
        alignment=TA_JUSTIFY,
        spaceAfter=12
    )

    # Add title
    elements.append(Paragraph(report.title, title_style))
    elements.append(Spacer(1, 12))

    # Add timestamp
    timestamp_text = f"Generated: {report.report_timestamp.strftime('%Y-%m-%d %H:%M:%S')}"
    elements.append(Paragraph(timestamp_text, styles['Normal']))
    elements.append(Spacer(1, 12))

    # Add executive summary
    elements.append(Paragraph("Executive Summary", heading_style))
    elements.append(Paragraph(report.executive_summary.replace('\n', '<br/>'), body_style))
    elements.append(Spacer(1, 12))

    # Add key takeaways
    if report.key_takeaways:
        elements.append(Paragraph("Key Takeaways", heading_style))
        for takeaway in report.key_takeaways:
            elements.append(Paragraph(f"• {takeaway}", body_style))
        elements.append(Spacer(1, 12))

    # Add main sections
    def add_section(section: ReportSection):
        elements.append(Paragraph(section.title, heading_style))
        # Convert markdown-like formatting to HTML for reportlab
        content = section.content.replace('\n', '<br/>')
        content = content.replace('**', '<b>').replace('*', '<i>')
        elements.append(Paragraph(content, body_style))

        if section.subsections:
            for subsection in section.subsections:
                add_section(subsection)

        elements.append(Spacer(1, 12))

    # Add all main sections
    add_section(report.introduction)
    add_section(report.methodology)
    elements.append(PageBreak())
    add_section(report.findings)
    add_section(report.analysis)
    add_section(report.insights)
    add_section(report.conclusions)

    # Add references
    if report.references:
        elements.append(PageBreak())
        elements.append(Paragraph("References", heading_style))
        for i, ref in enumerate(report.references, 1):
            ref_text = f"{i}. {ref.title} - {ref.url}"
            elements.append(Paragraph(ref_text, body_style))

    # Build PDF
    doc.build(elements)
    buffer.seek(0)
    return buffer.getvalue()


def generate_docx(report: ResearchReport) -> bytes:
    """Generate Word document from research report"""
    document = Document()

    # Add title
    title = document.add_heading(report.title, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add timestamp
    p = document.add_paragraph()
    p.add_run(f"Generated: {report.report_timestamp.strftime('%Y-%m-%d %H:%M:%S')}").italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add executive summary
    document.add_heading("Executive Summary", 1)
    document.add_paragraph(report.executive_summary)

    # Add key takeaways
    if report.key_takeaways:
        document.add_heading("Key Takeaways", 1)
        for takeaway in report.key_takeaways:
            document.add_paragraph(f"• {takeaway}", style='List Bullet')

    # Helper function to add sections
    def add_section(section: ReportSection, level: int = 1):
        document.add_heading(section.title, level)

        # Process content - handle markdown-like formatting
        content_lines = section.content.split('\n')
        for line in content_lines:
            if line.strip():
                p = document.add_paragraph()
                # Handle bold and italic
                parts = line.split('**')
                for i, part in enumerate(parts):
                    if i % 2 == 1:  # Bold parts
                        p.add_run(part).bold = True
                    else:
                        # Handle italics within non-bold parts
                        subparts = part.split('*')
                        for j, subpart in enumerate(subparts):
                            if j % 2 == 1:  # Italic parts
                                p.add_run(subpart).italic = True
                            else:
                                p.add_run(subpart)

        if section.subsections:
            for subsection in section.subsections:
                add_section(subsection, min(level + 1, 3))

    # Add main sections
    document.add_page_break()
    add_section(report.introduction)
    add_section(report.methodology)

    document.add_page_break()
    add_section(report.findings)
    add_section(report.analysis)
    add_section(report.insights)
    add_section(report.conclusions)

    # Add references
    if report.references:
        document.add_page_break()
        document.add_heading("References", 1)
        for i, ref in enumerate(report.references, 1):
            p = document.add_paragraph()
            p.add_run(f"{i}. {ref.title}").bold = True
            p.add_run(f" - Available at: {ref.url}")

    # Add word count at the end
    document.add_paragraph()
    p = document.add_paragraph()
    p.add_run(f"Total word count: {report.word_count}").italic = True

    # Save to bytes
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def generate_xml(report: ResearchReport) -> bytes:
    """Generate XML from research report"""
    # Convert report to dictionary
    report_dict = {
        'research_report': {
            'metadata': {
                'title': report.title,
                'generated': report.report_timestamp.isoformat(),
                'word_count': report.word_count
            },
            'executive_summary': report.executive_summary,
            'key_takeaways': {
                'takeaway': report.key_takeaways
            } if report.key_takeaways else {},
            'sections': {
                'introduction': section_to_dict(report.introduction),
                'methodology': section_to_dict(report.methodology),
                'findings': section_to_dict(report.findings),
                'analysis': section_to_dict(report.analysis),
                'insights': section_to_dict(report.insights),
                'conclusions': section_to_dict(report.conclusions)
            },
            'references': {
                'reference': [
                    {
                        'title': ref.title,
                        'url': ref.url,
                        'domain': ref.domain,
                        'type': ref.source_type
                    } for ref in report.references
                ]
            } if report.references else {}
        }
    }

    # Convert to XML
    xml = dicttoxml.dicttoxml(report_dict, custom_root='research_report', attr_type=False)
    return xml


def section_to_dict(section: ReportSection) -> Dict[str, Any]:
    """Convert ReportSection to dictionary for XML conversion"""
    section_dict = {
        'title': section.title,
        'content': section.content
    }

    if section.citations:
        section_dict['citations'] = {'citation': section.citations}

    if section.subsections:
        section_dict['subsections'] = {
            'subsection': [section_to_dict(sub) for sub in section.subsections]
        }

    return section_dict